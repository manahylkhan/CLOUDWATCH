"""Professional Word (.docx) report generator using python-docx."""
import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(hex_str: str):
    hex_str = hex_str.lstrip("#")
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _heading(doc: Document, text: str, level: int, color_hex: str = "#0A2342"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _rgb(color_hex)
    return h


def _add_colored_row(table, row_idx: int, bg_hex: str):
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex.lstrip("#"))
        tcPr.append(shd)


def generate_word_report(scan: dict, findings: list[dict], cis_results: list[dict] | None = None,
                         client_name: str | None = None, ai_roadmap: dict | None = None) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER ────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("CLOUDWATCH")
    run.font.size = Pt(36)
    run.font.color.rgb = _rgb("#E8651A")
    run.font.bold = True

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run("AWS Cloud Security Assessment Report").font.size = Pt(16)

    doc.add_paragraph()
    info = [
        ("Client", client_name or scan.get("account_alias", "AWS Account")),
        ("AWS Account ID", scan.get("account_id", "—")),
        ("Scan Date", datetime.fromisoformat(scan["started_at"]).strftime("%B %d, %Y") if scan.get("started_at") else "—"),
        ("Security Score", f"{scan.get('security_score', '—')}/100  —  Grade {scan.get('grade', '—')}"),
        ("Total Findings", str(scan.get("total_findings", 0))),
        ("Prepared By", "CLOUDWATCH AI Security Platform"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        _add_colored_row(table, i, "F3F4F6" if i % 2 == 0 else "FFFFFF")

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────────────
    _heading(doc, "Executive Summary", 1)

    score_para = doc.add_paragraph()
    score_para.add_run(f"Security Score: {scan.get('security_score','—')}/100  |  Grade: {scan.get('grade','—')}  |  "
                       f"Critical: {scan.get('critical_count',0)}  |  High: {scan.get('high_count',0)}  |  "
                       f"Medium: {scan.get('medium_count',0)}  |  Low: {scan.get('low_count',0)}").font.bold = True

    if scan.get("ai_summary"):
        doc.add_paragraph(scan["ai_summary"])

    # Findings by service table
    svc_counts: dict[str, dict] = {}
    for f in findings:
        svc = f.get("service", "Other")
        if svc not in svc_counts:
            svc_counts[svc] = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
        svc_counts[svc][f["severity"]] = svc_counts[svc].get(f["severity"], 0) + 1

    if svc_counts:
        doc.add_paragraph()
        t = doc.add_table(rows=len(svc_counts)+1, cols=5)
        t.style = "Table Grid"
        headers = ["Service", "Critical", "High", "Medium", "Low"]
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        _add_colored_row(t, 0, "0A2342")
        for j, cell in enumerate(t.rows[0].cells):
            cell.paragraphs[0].runs[0].font.color.rgb = _rgb("#FFFFFF")

        for row_i, (svc, counts) in enumerate(sorted(svc_counts.items()), start=1):
            row = t.rows[row_i]
            row.cells[0].text = svc
            row.cells[1].text = str(counts.get("Critical", 0))
            row.cells[2].text = str(counts.get("High", 0))
            row.cells[3].text = str(counts.get("Medium", 0))
            row.cells[4].text = str(counts.get("Low", 0))
            _add_colored_row(t, row_i, "F3F4F6" if row_i % 2 == 0 else "FFFFFF")

    doc.add_page_break()

    # ── FINDINGS BY SEVERITY ─────────────────────────────────────────────────
    sev_colors = {"Critical": "#DC2626", "High": "#EA580C", "Medium": "#D97706"}
    for severity in ("Critical", "High", "Medium", "Low"):
        sev_findings = [f for f in findings if f["severity"] == severity]
        if not sev_findings:
            continue

        _heading(doc, f"{severity} Findings ({len(sev_findings)})", 1, sev_colors.get(severity, "#374151"))

        for f in sev_findings:
            _heading(doc, f["title"], 2, sev_colors.get(severity, "#374151"))
            meta = doc.add_paragraph()
            meta.add_run(f"Service: {f.get('service','—')}  |  Resource: {f.get('resource_id','—')}  |  Region: {f.get('region','global')}").font.size = Pt(9)

            explanation = f.get("ai_explanation") or f.get("description") or ""
            if explanation:
                p = doc.add_paragraph()
                p.add_run("What This Means: ").font.bold = True
                p.add_run(explanation)

            if f.get("remediation"):
                p = doc.add_paragraph()
                p.add_run("How to Fix: ").font.bold = True
                p.add_run(f["remediation"])

            if f.get("remediation_effort"):
                doc.add_paragraph(f"Fix Effort: {f['remediation_effort']}")

        doc.add_page_break()

    # ── CIS BENCHMARK ───────────────────────────────────────────────────────
    if cis_results:
        _heading(doc, "CIS AWS Benchmark v2.0 Results", 1)
        passed = sum(1 for r in cis_results if r["status"] == "Pass")
        pct = round(passed / len(cis_results) * 100) if cis_results else 0
        doc.add_paragraph(f"Overall CIS Compliance: {pct}%  ({passed}/{len(cis_results)} checks passed)")

        t = doc.add_table(rows=len(cis_results)+1, cols=4)
        t.style = "Table Grid"
        for i, h in enumerate(["Check ID", "Title", "Status", "Evidence"]):
            t.rows[0].cells[i].text = h
            t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        _add_colored_row(t, 0, "0A2342")
        for cell in t.rows[0].cells:
            cell.paragraphs[0].runs[0].font.color.rgb = _rgb("#FFFFFF")

        for row_i, r in enumerate(sorted(cis_results, key=lambda x: x["cis_id"]), start=1):
            row = t.rows[row_i]
            row.cells[0].text = r["cis_id"]
            row.cells[1].text = r["title"]
            row.cells[2].text = r["status"]
            row.cells[3].text = (r.get("evidence", "") or "")[:100]
            bg = "D1FAE5" if r["status"] == "Pass" else "FEE2E2"
            _add_colored_row(t, row_i, bg)

        doc.add_page_break()

    # ── FIX ROADMAP ─────────────────────────────────────────────────────────
    if ai_roadmap:
        _heading(doc, "AI-Generated Fix Roadmap", 1)
        for label, items in [
            ("Today (Immediate)", ai_roadmap.get("immediate", [])),
            ("This Week", ai_roadmap.get("this_week", [])),
            ("This Month", ai_roadmap.get("this_month", [])),
        ]:
            if items:
                _heading(doc, label, 2)
                for item in items:
                    doc.add_paragraph(f"• {item}")
        doc.add_page_break()

    # ── APPENDIX ─────────────────────────────────────────────────────────────
    _heading(doc, "Appendix — All Findings", 1)
    t = doc.add_table(rows=len(findings)+1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Severity", "Service", "Resource", "Title"]):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    _add_colored_row(t, 0, "0A2342")
    for cell in t.rows[0].cells:
        cell.paragraphs[0].runs[0].font.color.rgb = _rgb("#FFFFFF")

    sorted_f = sorted(findings, key=lambda x: ["Critical","High","Medium","Low","Info"].index(x["severity"]))
    for row_i, f in enumerate(sorted_f, start=1):
        row = t.rows[row_i]
        row.cells[0].text = f["severity"]
        row.cells[1].text = f.get("service", "—")
        row.cells[2].text = (f.get("resource_id") or "—")[:30]
        row.cells[3].text = f["title"][:80]
        _add_colored_row(t, row_i, "F3F4F6" if row_i % 2 == 0 else "FFFFFF")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
